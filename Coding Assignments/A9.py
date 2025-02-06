from dotenv import load_dotenv
load_dotenv()

import os
GOOGLE_API_KEY = os.environ['GOOGLE_API_KEY']


import streamlit as st


def extract_texts(file):
    def _extract_pdf(file):
        from pypdf import PdfReader

        reader = PdfReader(file)
        texts = ""
        for page in reader.pages:
            texts += page.extract_text()
        return texts
    
    def _extract_docx(file):
        from docx import Document

        doc = Document(file)
        texts = []
        for para in doc.paragraphs:
            texts.append(para.text)
        return '\n\n'.join(texts)
    
    def _extract_csv(file):
        from csv import reader
        
        texts = []
        with open(file, 'r') as f:
            csv_file = reader(f)
            for row in csv_file:
                texts.append(', '.join(row))
        return '\n'.join(texts)

    
    texts = ""
    ext = os.path.splitext(file.name)[-1]
    match ext:
        case '.pdf':
            texts += _extract_pdf(file)
        case '.docx':
            texts += _extract_docx(file)
        case ".csv":
            texts += _extract_csv(file)

    return texts

def make_chunks(texts, file):
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    from langchain_core.documents import Document

    
    splitter = RecursiveCharacterTextSplitter(separators=['\n\n', '\n', ' '], chunk_size=750, chunk_overlap=100)
    docs = []
    for chunk in splitter.split_text(texts):
        docs.append(Document(page_content=chunk, metadata={"source": file}))
    return docs

def make_vecdb(chunks):
    from langchain_community.vectorstores import FAISS
    from langchain_huggingface.embeddings import HuggingFaceEmbeddings

    return FAISS.from_documents(chunks, HuggingFaceEmbeddings(model_name='all-MiniLM-L6-v2'))

def make_chain(vecdb):
    from langchain.chains.history_aware_retriever import create_history_aware_retriever
    from langchain.chains.combine_documents import create_stuff_documents_chain
    from langchain.chains.retrieval import create_retrieval_chain
    from langchain.prompts import ChatPromptTemplate, MessagesPlaceholder
    from langchain_google_genai import ChatGoogleGenerativeAI

    rephrase_prompt = ChatPromptTemplate.from_messages([
        ("system", "You are an URDU Chatbot, so uou must ALWAYS answer anything in URDU. You may be given some context and chat history, so use them to answer the question at the end."),
        MessagesPlaceholder(variable_name="chat_history"),
        ("user", "{input}")
    ])
    qa_prompt = ChatPromptTemplate.from_messages([
        ("system", "Context: {context}"),
        ("user", "Question: {input}\nAnswer in URDU:")
    ])

    llm = ChatGoogleGenerativeAI(model="gemini-1.5-flash")

    retriever = create_history_aware_retriever(llm, vecdb.as_retriever(), rephrase_prompt)
    chain = create_retrieval_chain(retriever, create_stuff_documents_chain(llm, qa_prompt))
    return chain

def handle_query(query):
    from langchain_core.messages import HumanMessage, AIMessage
    from streamlit_chat import message

    history = st.session_state['chat_history']
    chain = st.session_state['chain']
    
    resp = chain.invoke({"input": query, "chat_history": history})
    st.session_state['chat_history'].extend([
        HumanMessage(content=query),
        AIMessage(content=resp['answer']),
    ])

    with st.container():
        for i, msg in enumerate(st.session_state['chat_history']):
            message(msg.content, is_user=(i % 2 == 0), key=str(i))
    
    return resp['answer']

def handle_audio(audio):
    from scipy.io import wavfile
    import numpy as np


    def _speech_to_text(audio):
        from transformers import pipeline
        
        pipe = pipeline("automatic-speech-recognition", model="Talha/URDU-ASR")
        text = pipe(audio)
        text = text['text']
        return text

    def _text_to_speech(text):
        from transformers import VitsModel, AutoTokenizer
        import torch

        model = VitsModel.from_pretrained("facebook/mms-tts-urd-script_arabic")
        tok = AutoTokenizer.from_pretrained("facebook/mms-tts-urd-script_arabic")

        inputs = tok(text, return_tensors='pt')
        with torch.no_grad():
            audio = model(**inputs).waveform
        audio = audio.numpy()
        audio = np.squeeze(audio)
        wavfile.write("out.wav", model.config.sampling_rate, audio)
        return "out.wav"

    
    with open("tmp.wav", "wb") as f:
        f.write(audio.getvalue())
    
    # test with test.wav instead of user's audio input
    # audio = wavfile.read("tmp.wav")
    audio = wavfile.read("test.wav")
    if os.path.exists("tmp.wav"):
        os.remove("tmp.wav")

    audio = np.array(audio[1], dtype=np.float32)

    text = _speech_to_text(audio)
    response = handle_query(text)
    audio_path = _text_to_speech(response)
    
    st.audio(audio_path)
    if os.path.exists(audio_path):
        os.remove(audio_path)
    

def main():
    st.set_page_config(page_title="URDU Chatbot")
    st.header("URDU Chatbot")

    for key, default in [("chain", None), ("chat_history", []), ("ready", False)]:
        if key not in st.session_state:
            st.session_state[key] = default

    with st.sidebar:
        uploaded = st.file_uploader("Choose Files", ["pdf", "csv", "docx"], accept_multiple_files=True)
        if uploaded:
            if not GOOGLE_API_KEY:
                st.info("Please specify LLM API key")
                st.stop()
            
            all_chunks = []
            for file in uploaded:
                st.write(f"Loading File {file.name}...")
                texts = extract_texts(file)
                st.write(f"Chunking File {file.name}...")
                chunks = make_chunks(texts, file.name)
                st.write(f"Loaded File {file.name}!")
                all_chunks.extend(chunks)
            
            vecDB = make_vecdb(chunks)
            
            st.session_state['chain'] = make_chain(vecDB)
            st.session_state['ready'] = True
        
    if st.session_state['ready']:
        query = st.chat_input("Chat with Documents")
        audio = st.audio_input("Talk with Documents")

        if query:
            handle_query(query)
        
        if audio:
            handle_audio(audio)

if __name__ == "__main__":
    main()
